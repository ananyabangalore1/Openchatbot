import openai
import streamlit as st
#from creds import openai_api_key
import docx
import os
from dotenv import load_dotenv


load_dotenv('key.env')

# Get the API key from the environment variable
api_key = os.getenv("OPENAI_API_KEY")

if api_key:
    openai.api_key = api_key
else:
    raise ValueError("API key not found. Please check your key.env file.")

#def extract_text_from_docx(file):
 #   doc = Document(file)
   # text = []
    #for para in doc.paragraphs:
     #   text.append(para.text)
    #return "\n".join(text)

# Streamlit App UI

# Custom CSS for colors
st.markdown(
    """
    <style>
    .stApp {
        background-color: #c4ae78; /* Cream color */
    }
    /* Sidebar styling */
    .css-1d391kg {
        background-color: #171515; /* Brown color */
        color: white;
    }
    .css-1d391kg p, .css-1d391kg h1, .css-1d391kg h2, .css-1d391kg h3 {
        color: #c4ae78;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("My Bot")

# Initialize session state variables
if "messages" not in st.session_state:
    st.session_state.messages = []

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# Sidebar for previous chats
with st.sidebar:
    st.header("Previous Chats")
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history):
            st.write(f"Chat {i + 1}:")
            for message in chat:
                st.markdown(f"**{message['role'].capitalize()}:** {message['content']}")
            st.write("---")  # Separator between chats
    else:
        st.write("No previous conversations.")

    # Button to start a new chat
    if st.button("Start New Chat"):
        if st.session_state.messages:
            # Append the current conversation to chat history
            st.session_state.chat_history.append(st.session_state.messages.copy())
        # Reset messages for a new conversation
        st.session_state.messages = []

# Function to read few-shot examples from a docx file
def load_few_shot_examples(docx_file_path):
    # Load the .docx file
    doc = docx.Document(docx_file_path)
    few_shot_prompt = ""

    # Read paragraphs from the .docx file and append them to the prompt
    for paragraph in doc.paragraphs:
        few_shot_prompt += paragraph.text + "\n"

    return few_shot_prompt

# Path to the document with few-shot examples
# Assuming your document is located in a 'data' subfolder within the project directory
current_directory = os.path.dirname(__file__)  # Path to the script's location
docx_file_name = "few_shot_examples.docx"
docx_file_path = os.path.join(current_directory, "data", docx_file_name)

# Load the few-shot examples from the document
few_shot_prompt = load_few_shot_examples(docx_file_path)

# Display current chat messages with icons
for message in st.session_state.messages:
    if message['role'] == 'user':
        st.markdown(f"**👤 User:** {message['content']}")  # User icon
    else:
        st.markdown(f"**🤖 Assistant:** {message['content']}")  # Assistant icon

# Text input for user to send messages
prompt = st.text_input("What's up?", placeholder="Type your message here...")

if prompt:
    # Combine the few-shot examples with the user's input
    full_prompt = few_shot_prompt + "\n" + prompt

    # Store user message in session state
    st.session_state.messages.append({"role": "user", "content": prompt})

    # Display user's message with an icon
    st.markdown(f"**👤 User:** {prompt}")

    try:
        # Call the OpenAI API to get a response using few-shot learning
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": few_shot_prompt}, {"role": "user", "content": prompt}],
            max_tokens=150
        )

        # Extract assistant's response
        assistant_response = response['choices'][0]['message']['content']

        # Display assistant's response with an icon
        st.markdown(f"**🤖 Assistant:** {assistant_response}")

        # Store assistant's response in session state
        st.session_state.messages.append({"role": "assistant", "content": assistant_response})

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

# Button to clear chat history
if st.sidebar.button("Clear Chat History"):
    st.session_state.chat_history = []
    st.session_state.messages = []

